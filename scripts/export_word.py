import argparse
import os
from docx import Document
from docx.shared import Pt, Inches

def md_to_docx(input_path, output_path, title):
    doc = Document()
    head = doc.add_heading(title, 0)
    head.alignment = 1 # Center
    
    with open(input_path, "r", encoding="utf-8") as f:
        for line in f:
            line = line.strip()
            if not line:
                doc.add_paragraph("")
                continue
            if line.startswith("# "):
                doc.add_heading(line[2:].strip(), level=1)
            elif line.startswith("## "):
                doc.add_heading(line[3:].strip(), level=2)
            elif line.startswith("### "):
                doc.add_heading(line[4:].strip(), level=3)
            elif line.startswith("- ") or line.startswith("* "):
                p = doc.add_paragraph(line[2:].strip(), style='List Bullet')
            elif line.startswith("> ") or line.startswith("  "):
                p = doc.add_paragraph(line[2:].strip(), style='Quote')
            else:
                doc.add_paragraph(line)
                
    doc.save(output_path)
    print(f"Word document saved to {output_path}")

def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--type", required=True, choices=["exercises", "report"])
    parser.add_argument("--input", required=True, help="Input markdown text file")
    parser.add_argument("--output", required=True, help="Output docx file path")
    parser.add_argument("--title", default="辅导文档")
    args = parser.parse_args()
    
    if args.type in ["exercises", "report"]:
        md_to_docx(args.input, args.output, args.title)

if __name__ == "__main__":
    main()
